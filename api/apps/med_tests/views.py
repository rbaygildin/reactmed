from json import loads

from io import StringIO, BytesIO
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.urls import reverse
from django.views.decorators.csrf import csrf_exempt
from docx import Document

from apps.core.models import Patient, TestRec


@login_required
@csrf_exempt
def create_action(request, patient_id):
    if request.method == 'GET':
        return render(request, 'med_tests/create.html', {'patient_id': patient_id})
    elif request.method == 'POST':
        patient = Patient.objects.get(pk=patient_id)
        d = loads(request.body.decode('utf-8'))
        TestRec.objects.create(
            name=d.get('name'),
            short_name=d.get('short_name'),
            description=d.get('description', None),
            real_inds=d.get('real_inds', None),
            int_inds=d.get('int_inds', None),
            text_inds=d.get('text_inds', None),
            info=d.get('info'),
            summary=d.get('summary'),
            test_date=d.get('test_date'),
            patient=patient
        )
    return redirect(reverse('patients:show', kwargs={'patient_id': patient_id}))


@login_required
def show_action(request, test_id):
    test_rec = TestRec.objects.get(pk=test_id)

    patient = test_rec.patient
    return render(
        request,
        'med_tests/show.html',
        {
            'test_rec': test_rec,
            'patient': patient
        }
    )


@login_required
def delete_action(request, test_id):
    test_rec = TestRec.objects.get(pk=test_id)
    test_rec.delete()
    return redirect(reverse('patients:list'))


@login_required
def dynamics_action(request, patient_id=None):
    return render(
        request,
        'med_tests/dynamics.html',
        {
            'patient': Patient.objects.get(pk=patient_id)
        }
    )


def val2str(val):
    if val is None:
        return ''
    return str(val)


@login_required
def report_action(request, test_id):
    test_rec = TestRec.objects.get(pk=test_id)
    patient = test_rec.patient
    document = Document()
    document.add_heading(test_rec.name, 0)
    patient_table = document.add_table(rows=1, cols=2)
    # Patient full name
    row = patient_table.add_row().cells
    row[0].text = 'ФИО'
    row[1].text = patient.full_name
    # Gender
    row = patient_table.add_row().cells
    row[0].text = 'Пол'
    row[1].text = patient.gender
    # Age
    row = patient_table.add_row().cells
    row[0].text = 'Возраст'
    row[1].text = str(patient.age)
    # Test date
    row = patient_table.add_row().cells
    row[0].text = 'Дата проведения анализа'
    row[1].text = str(test_rec.test_date)
    if test_rec.real_inds is not None and len(test_rec.real_inds.keys()) != 0:
        document.add_heading('Вещественные показатели', level=1)
        table = document.add_table(rows=1, cols=4)
        int_inds = table.rows[0].cells
        int_inds[0].text = 'Показатель'
        int_inds[1].text = 'Значение'
        int_inds[2].text = 'Норма'
        int_inds[3].text = 'Единица измерения'
        for ind, ind_value in test_rec.real_inds.items():
            row_cells = table.add_row().cells
            row_cells[0].text = ind_value.get('name')
            row_cells[1].text = str(ind_value.get('value'))
            row_cells[2].text = '%s - %s' % (
            val2str(ind_value.get('min_norm', '')), val2str(ind_value.get('max_norm', '')))
            row_cells[3].text = val2str(ind_value.get('unit', ' - '))
    if test_rec.int_inds is not None and len(test_rec.int_inds.keys()) != 0:
        document.add_heading('Целые показатели', level=1)
        table = document.add_table(rows=1, cols=4)
        int_inds = table.rows[0].cells
        int_inds[0].text = 'Показатель'
        int_inds[1].text = 'Значение'
        int_inds[2].text = 'Норма'
        int_inds[3].text = 'Единица измерения'
        for ind, ind_value in test_rec.int_inds.items():
            row_cells = table.add_row().cells
            row_cells[0].text = ind_value.get('name')
            row_cells[1].text = str(ind_value.get('value'))
            row_cells[2].text = '%s - %s' % (
            val2str(ind_value.get('min_norm', '')), val2str(ind_value.get('max_norm', '')))
            row_cells[3].text = val2str(ind_value.get('unit', ' - '))
    if test_rec.text_inds is not None and len(test_rec.text_inds.keys()) != 0:
        document.add_heading('Строковые показатели', level=1)
        table = document.add_table(rows=1, cols=2)
        int_inds = table.rows[0].cells
        int_inds[0].text = 'Показатель'
        int_inds[1].text = 'Значение'
        for ind, ind_value in test_rec.text_inds.items():
            row_cells = table.add_row().cells
            row_cells[0].text = ind_value.get('name')
            row_cells[1].text = ind_value.get('value')
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'inline; filename="%s - %s (%s)"' % (
        patient.short_name, test_rec.name, test_rec.test_date
    )
    response['Content-Length'] = length
    return response
